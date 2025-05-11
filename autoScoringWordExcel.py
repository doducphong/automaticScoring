# import pytesseract
# import torch
# from PIL import Image
# from docx import Document
# from docx.shared import RGBColor
# from transformers import LayoutLMTokenizer, LayoutLMForTokenClassification, ViTFeatureExtractor, ViTForImageClassification

# # Load mô hình LayoutLM để kiểm tra bố cục
# tokenizer = LayoutLMTokenizer.from_pretrained("microsoft/layoutlm-base-uncased")
# layout_model = LayoutLMForTokenClassification.from_pretrained("microsoft/layoutlm-base-uncased")

# # Load mô hình ViT để kiểm tra lỗi hình ảnh
# feature_extractor = ViTFeatureExtractor.from_pretrained("google/vit-base-patch16-224")
# vit_model = ViTForImageClassification.from_pretrained("google/vit-base-patch16-224")


# def extract_text_from_docx(doc_path):
#     """ Trích xuất văn bản từ file .docx """
#     doc = Document(doc_path)
#     text_data = []
#     for para in doc.paragraphs:
#         text_data.append(para.text)
#     return text_data


# def detect_text_errors(text_data):
#     """ Machine Learning kiểm tra lỗi font chữ, căn lề, tiêu đề """
#     errors = []
#     for i, text in enumerate(text_data):
#         if len(text) > 0 and text.isupper():  # Ví dụ: Kiểm tra nếu text viết hoa toàn bộ
#             errors.append((i, "Lỗi tiêu đề: Viết hoa toàn bộ"))
#         if "  " in text:  # Kiểm tra khoảng trắng dư thừa
#             errors.append((i, "Lỗi căn lề: Có khoảng trắng dư"))
#     return errors


# def detect_layout_errors(text_data):
#     """ Deep Learning với LayoutLM để kiểm tra lỗi bố cục """
#     inputs = tokenizer(text_data, return_tensors="pt", padding=True, truncation=True)
#     outputs = layout_model(**inputs)
#     logits = outputs.logits
#     predictions = torch.argmax(logits, dim=-1)

#     errors = []
#     for i, pred in enumerate(predictions[0].tolist()):
#         if pred != 0:  # Nếu model dự đoán có lỗi
#             errors.append((i, "Lỗi bố cục: Heading không đúng định dạng"))
#     return errors


# def extract_images_from_docx(doc_path):
#     """ Trích xuất hình ảnh từ file .docx """
#     doc = Document(doc_path)
#     image_paths = []
#     for i, rel in enumerate(doc.part.rels):
#         if "image" in doc.part.rels[rel].target_ref:
#             image_path = f"image_{i}.png"
#             with open(image_path, "wb") as f:
#                 f.write(doc.part.rels[rel].target_part.blob)
#             image_paths.append(image_path)
#     return image_paths


# def detect_image_errors(image_paths):
#     """ Deep Learning với ViT để kiểm tra lỗi hình ảnh """
#     errors = []
#     for image_path in image_paths:
#         image = Image.open(image_path)
#         inputs = feature_extractor(images=image, return_tensors="pt")
#         outputs = vit_model(**inputs)
#         preds = torch.argmax(outputs.logits, dim=-1)
#         if preds != 0:  # Nếu model phát hiện lỗi
#             errors.append((image_path, "Lỗi hình ảnh: Sai kích thước hoặc bố cục"))
#     return errors


# def highlight_errors(doc_path, errors):
#     """ Tô màu đỏ lỗi trong file .docx """
#     doc = Document(doc_path)
#     for index, error_msg in errors:
#         if isinstance(index, int):  # Lỗi văn bản
#             para = doc.paragraphs[index]
#             run = para.add_run(f"  (⚠ {error_msg})")
#             run.font.color.rgb = RGBColor(255, 0, 0)  # Đánh dấu lỗi màu đỏ
#         else:  # Lỗi hình ảnh
#             doc.add_paragraph(f"⚠ {error_msg} tại {index}").bold = True

#     output_path = "output_with_errors.docx"
#     doc.save(output_path)
#     print(f"✅ Báo cáo lỗi đã lưu tại: {output_path}")


# def main(doc_path):
#     print("📂 Đang xử lý file:", doc_path)

#     # Bước 1: Trích xuất văn bản
#     text_data = extract_text_from_docx(doc_path)
    
#     # Bước 2: Kiểm tra lỗi văn bản (ML + DL)
#     text_errors = detect_text_errors(text_data)
#     layout_errors = detect_layout_errors(text_data)

#     # Bước 3: Kiểm tra lỗi hình ảnh
#     image_paths = extract_images_from_docx(doc_path)
#     image_errors = detect_image_errors(image_paths)

#     # Bước 4: Ghi nhận lỗi
#     all_errors = text_errors + layout_errors + image_errors
#     if all_errors:
#         print("❌ Phát hiện lỗi trong tài liệu!")
#         for error in all_errors:
#             print(f"- {error[1]} (Dòng {error[0]})")
#         highlight_errors(doc_path, all_errors)
#     else:
#         print("✅ Tài liệu không có lỗi!")

    
# # Chạy kiểm tra trên file .docx
# file_path = "D:\doAnTuChamDiemFileWordExcel\test.docx"  # Đổi thành file của bạn
# main(file_path)



# from docx import Document

# doc = Document(r"D:\doAnTuChamDiemFileWordExcel\test.docx")
# for para in doc.paragraphs:
#     para_style = para.style
#     for run in para.runs:
#         font_name = run.font.name if run.font.name else para_style.font.name
#         font_size = run.font.size if run.font.size else para_style.font.size
#         print(f"Font: {font_name}, Size: {font_size.pt}, style: {para.style.name}")

# sections = doc.sections
# for section in sections:
#         print(f"Top: {section.top_margin.cm}, Bottom: {section.bottom_margin.cm}")
#         print(f"Left: {section.left_margin.cm}, Right: {section.right_margin.cm}")

# from transformers import LayoutLMForSequenceClassification, LayoutLMTokenizer

# tokenizer = LayoutLMTokenizer.from_pretrained("microsoft/layoutlm-base-uncased")
# model = LayoutLMForSequenceClassification.from_pretrained("microsoft/layoutlm-base-uncased")

# inputs = tokenizer("Tiêu đề: Báo cáo tài chính", return_tensors="pt")
# outputs = model(**inputs)
# print(outputs.logits)

import zipfile
import xml.etree.ElementTree as ET
from docx import Document
from docx.oxml.ns import qn
import sys
sys.stdout.reconfigure(encoding='utf-8')
import base64
import io
import re

# Định nghĩa namespace cần dùng
namespaces = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
}


def has_section_break(paragraph):
    """
    Kiểm tra xem đoạn paragraph có chứa section break hay không thông qua tag <w:sectPr> trong pPr.
    """
    # Dùng xpath để kiểm tra nếu tồn tại phần tử sectPr trong paragraph's properties
    sectPrs = paragraph._p.xpath('./w:pPr/w:sectPr')
    return len(sectPrs) > 0

def extract_sections(docx_path):
    doc = Document(docx_path)
    
    sections_content = []  # Danh sách chứa nội dung từng section (danh sách các đoạn văn)
    current_section = []   # Danh sách tạm chứa text của section hiện tại

    # Duyệt qua từng paragraph trong tài liệu
    for para in doc.paragraphs:
        # Thêm nội dung paragraph vào section hiện tại
        current_section.append(para.text)
        
        # Nếu phát hiện section break (có tag <w:sectPr> trong pPr của paragraph)
        if has_section_break(para):
            # Lưu lại section hiện tại
            sections_content.append(current_section)
            # Tạo section mới cho phần sau
            current_section = []
    
    # Sau khi duyệt hết, nếu còn nội dung chưa lưu (phần cuối không có section break explicit)
    if current_section:
        sections_content.append(current_section)
    
    return sections_content

def remove_namespace(tag):
    return tag.split('}')[-1] if '}' in tag else tag

def find_wordart_texts(element):
    tag = remove_namespace(element.tag)

    if tag == "shape":
        print(f"[DEBUG] Found <shape> tag with attributes: {element.attrib}")
        for attr, val in element.attrib.items():
            if "gfxdata" in attr.lower():
                print("[DEBUG] Found gfxdata attribute, trying to decode...")
                try:
                    decoded = base64.b64decode(val)
                    with zipfile.ZipFile(io.BytesIO(decoded)) as inner_zip:
                        print(f"[DEBUG] Inner zip file contents: {inner_zip.namelist()}")
                        for name in inner_zip.namelist():
                            if name.endswith(".xml"):
                                inner_xml = inner_zip.read(name).decode("utf-8", errors="ignore")
                                matches = re.findall(r'string="(.*?)"', inner_xml)
                                for text in matches:
                                    print(f"[FOUND WordArt] Văn bản: {text} (trong {name})")
                except zipfile.BadZipFile:
                    print("[ERROR] gfxdata is not a valid zip file")
    # Duyệt tiếp các node con
    for child in element:
        find_wordart_texts(child)

def check_two_columns(docx_path):
    # Mở file DOCX dưới dạng file ZIP
    with zipfile.ZipFile(docx_path) as docx_zip:
        # Đọc nội dung XML của file document.xml
        xml_content = docx_zip.read("word/document.xml")
    
    # Phân tích cú pháp XML
    tree = ET.fromstring(xml_content)
    root = ET.fromstring(xml_content)
    find_wordart_texts(root)
    # Iterate through each child of the root element
    
    
    # Tìm tất cả các section (w:sectPr)
    # Lưu ý: Các section thường xuất hiện ở cuối document hoặc giữa các paragraph
    # for sect in tree.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr"):
    #     cols = sect.find("w:cols", namespaces)
    #     if cols is not None:
    #         num_cols = cols.attrib.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num")
    #         if num_cols == "2":
    #             print("Section 2 col.")
    #         else:
    #             print(f"Section 1 col: {num_cols}")
    #     else:
    #         print("Section not col.")
    for para in tree.findall(".//w:p", namespaces):
    # Lấy phần tử pPr (paragraph properties)
        pPr = para.find("w:pPr", namespaces)
        if pPr is not None:
            drop_cap = pPr.find("w:dropCap", namespaces)
            if drop_cap is not None:
                print("Drop cap found in paragraph! Type:")
                drop_val = drop_cap.attrib.get(f"{{{namespaces['w']}}}val", "none")
                if drop_val in ("drop", "margin"):
                    print("Drop cap found in paragraph! Type:", drop_val)

# Sử dụng hàm trên với file Word của bạn
docx_path = r"D:\doAnTuChamDiemFileWordExcel\dropcap.docx"
# sections = extract_sections(docx_path)

# In nội dung từng section ra
# for idx, section in enumerate(sections, start=1):
#     print(f"--- Section {idx} ---")
#     for para in section:
#         print(para)
#     print("\n")
check_two_columns(docx_path)